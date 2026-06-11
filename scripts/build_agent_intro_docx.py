"""
build_agent_intro_docx.py
FA 자동화팀 AI 에이전트 소개 Word 문서 생성
출력: outputs/에이전트_소개_2026-05-29.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "에이전트_소개_2026-05-29.docx")

# ── 색상 상수 ──────────────────────────────────────────────
COLOR_TITLE   = RGBColor(0x00, 0x33, 0x99)
COLOR_HEADING = RGBColor(0x00, 0x33, 0x99)
COLOR_ACCENT  = RGBColor(0x00, 0x00, 0xFF)
COLOR_GREEN   = RGBColor(0x00, 0x66, 0x00)
COLOR_GRAY    = RGBColor(0x60, 0x60, 0x60)
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_HEADER_BG = RGBColor(0x00, 0x33, 0x99)
COLOR_ROW_ODD   = RGBColor(0xF0, 0xF4, 0xFF)
COLOR_ROW_EVEN  = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_SUB_HDR   = RGBColor(0xCC, 0xD9, 0xFF)


# ── 헬퍼 ──────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_font(doc, name_kr="맑은 고딕", name_en="Arial Narrow", size_pt=10):
    style = doc.styles["Normal"]
    style.font.name = name_en
    style.font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_kr)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_note(doc, text):
    pn = doc.add_paragraph(text)
    pn.paragraph_format.space_before = Pt(4)
    for run in pn.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GREEN
        run.bold = True


def make_table(doc, headers, rows, col_widths, header_bg=COLOR_HEADER_BG):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        if header_bg == COLOR_HEADER_BG:
            r.font.color.rgb = COLOR_WHITE

    for idx, row_data in enumerate(rows):
        bg = COLOR_ROW_ODD if idx % 2 == 0 else COLOR_ROW_EVEN
        row = tbl.add_row()
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if ci == 0:
                r.bold = True
                if header_bg == COLOR_HEADER_BG:
                    r.font.color.rgb = COLOR_ACCENT

    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths):
                cell.width = col_widths[i]
    return tbl


# ── 표지 ──────────────────────────────────────────────────

def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("FA 자동화팀 AI 에이전트 소개")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = COLOR_TITLE

    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Agent Team 구성 및 협업 패턴 안내")
    r2.font.size = Pt(14)
    r2.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_date.add_run("2026-05-29")
    r3.font.size = Pt(12)
    r3.font.color.rgb = COLOR_GRAY

    p_note = doc.add_paragraph()
    p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p_note.add_run("(출처: CLAUDE.md · .claude/agents/ 에이전트 정의 파일)")
    r4.font.size = Pt(9)
    r4.font.color.rgb = COLOR_GREEN
    r4.bold = True

    doc.add_page_break()


# ── 개요 표 ───────────────────────────────────────────────

def build_overview(doc):
    add_heading(doc, "1. 에이전트 한눈에 보기", level=1, color=COLOR_HEADING)

    p = doc.add_paragraph(
        "본 팀은 5개의 AI 에이전트로 구성되며, 각 에이전트는 도메인별 역할을 분담하여 "
        "병렬·협업 방식으로 보고서를 작성합니다."
    )
    p.paragraph_format.space_after = Pt(8)

    headers = ["에이전트", "역할 요약", "주요 도구", "사용 시점"]
    rows = [
        ("data-teammate",          "KPI·Capex·MRM·DST 데이터 조회",           "Read, Grep, Glob, Bash",                     "KPI 진척·투자비·양산·R&D 데이터"),
        ("ops-teammate",           "협력사·HR·결재·사내문서 검색",              "Read, Grep, Glob, Bash",                     "협력사 현황·HR·승인·사내 PDF/PPT"),
        ("tech-research-teammate", "외부 기술동향·경쟁사·기술평가",              "Read, Grep, Glob, Bash, WebSearch, WebFetch","신규 과제 발굴·KPI 재설정"),
        ("document-writer",        "markdown(.md) 양식 빈칸 채워 문서 완성",   "Read, Write, Edit, Glob, Grep",              "회의록·주간보고·월간보고 작성"),
        ("ppt-writer",             "PPTX 양식 분석·빈칸 채워 PPT 생성",        "Read, Write, Edit, Glob, Grep, Bash",        "주간보고·임원보고 PPT 작성"),
    ]
    col_widths = [Cm(3.8), Cm(5.2), Cm(5.0), Cm(4.0)]
    make_table(doc, headers, rows, col_widths)

    doc.add_paragraph()
    add_note(doc, "(출처: CLAUDE.md 에이전트 5명 항목)")
    doc.add_page_break()


# ── 에이전트 상세 ─────────────────────────────────────────

AGENTS_DETAIL = [
    {
        "title": "2-1. data-teammate — 사내 정량 데이터 조회",
        "role": "KPI·Capex(투자)·MRM(양산 일정)·DST(R&D 진척) 등 숫자와 일정 데이터를 sample_data/ 및 references/ 에서 조회·집계합니다.",
        "when": "KPI 진척률, 투자비, 양산 로드맵, R&D 과제 진척을 확인해야 할 때",
        "tools": "Read, Grep, Glob, Bash",
        "areas": [
            ("KPI·Roadmap",     "분기별 추이, 목표 달성률 집계"),
            ("Capex·면적·구매",  "투자비 예산, 발주 현황, 면적 사용률"),
            ("MRM (양산 로드맵)", "양산 일정, 기술과제 진척 현황"),
            ("DST (R&D 미래기술)", "미래기술 과제 진척 및 마일스톤"),
        ],
        "location": "sample_data/ 폴더, references/roadmap/, references/policy/",
        "note": "(출처: CLAUDE.md [data-teammate 항목], .claude/agents/data-teammate.md)",
    },
    {
        "title": "2-2. ops-teammate — 사내 운영 문서 검색",
        "role": "협력사·HR(출장/교육/포상)·결재 큐·사내문서를 검색하고 현황을 정리합니다.",
        "when": "협력사 현황, HR 활동(출장·교육·포상), 승인 대기 건, 사내 PDF/PPT 검색이 필요할 때",
        "tools": "Read, Grep, Glob, Bash",
        "areas": [
            ("협력사 (Partners)",  "협력사 풀, 벤더 등록 현황, 평가 결과"),
            ("HR (출장·교육·포상)", "출장 신청, 교육 이수, 포상 추천 내역"),
            ("결재 큐 (Approvals)", "승인 대기·승인·반려 처리 현황"),
            ("사내 문서",           "PDF/PPT/텍스트 전문 검색"),
        ],
        "location": "sample_data/, references/policy/, references/과거회의록/",
        "note": "(출처: CLAUDE.md [ops-teammate 항목], .claude/agents/ops-teammate.md)",
    },
    {
        "title": "2-3. tech-research-teammate — 외부 기술 리서치",
        "role": "AMR/협동로봇/디지털트윈/스마트물류/AI/피지컬AI 등 배터리 공장 FA 영역의 외부 기술 동향·경쟁사 정보·기술 평가를 수집하고 아카이빙합니다.",
        "when": "신규 과제 발굴, KPI 재설정, 로드맵 검토, 경쟁사 분석이 필요할 때",
        "tools": "Read, Grep, Glob, Bash, Write, WebSearch, WebFetch",
        "areas": [
            ("외부 기술 동향", "글로벌(영어) 및 한국어 뉴스·기술 블로그·학술 자료 수집"),
            ("경쟁사 정보",   "CATL, 삼성SDI, BYD, Tesla, Northvolt 등 동향 추적"),
            ("기술 평가",     "TRL 산정, LGES 적용 가능성, 리스크 분석"),
        ],
        "location": "수집 결과 자동 아카이빙: references/기술자료/ 또는 references/경쟁사/",
        "note": "(출처: CLAUDE.md [tech-research-teammate 항목], .claude/agents/tech-research-teammate.md)",
    },
    {
        "title": "2-4. document-writer — Markdown 보고서 작성",
        "role": "templates/ 의 .md 양식을 읽어 {{ placeholder }} 를 채우고 outputs/ 에 완성된 문서를 저장합니다.",
        "when": '"회의록 양식으로 정리해줘", "주간보고 작성해줘" 같은 요청',
        "tools": "Read, Write, Edit, Glob, Grep",
        "areas": [
            ("회의록",   "templates/회의록.md 기반"),
            ("주간보고", "templates/주간보고.md 기반"),
            ("월간보고", "templates/월간보고.md 기반"),
        ],
        "location": "templates/ (양식), references/ (참조), outputs/ (결과물 저장)",
        "note": "특이 사항: 녹취록 입력 시 meeting-minutes 스킬로 자동 위임\n(출처: .claude/agents/document-writer.md)",
    },
    {
        "title": "2-5. ppt-writer — PPTX 보고서 작성",
        "role": "PPT(.pptx) 양식을 분석해 빈칸을 추론하고, 다른 팀원에게 데이터를 요청해 채운 PPTX 파일을 생성합니다.",
        "when": '"주간보고 PPT 작성해줘", "이 PPT 양식 채워줘" 같은 요청',
        "tools": "Read, Write, Edit, Glob, Grep, Bash",
        "areas": [
            ("구조 추출", "scripts/ppt_extract.py 로 양식 슬라이드 구조 분석"),
            ("빈칸 추론", "placeholder 목록 도출 및 필요 데이터 식별"),
            ("데이터 수집", "data-teammate / ops-teammate 에 데이터 요청"),
            ("PPT 채움", "scripts/ppt_fill.py 로 placeholder 치환·저장"),
        ],
        "location": "templates/ (양식), outputs/<이름>_YYYY-MM-DD.pptx (결과물)",
        "note": "LGES PPT 작업 가이드(templates/사내양식/LGES_PPT_작업_가이드.md) 및 분류형 레이아웃 표준 준수\n(출처: .claude/agents/ppt-writer.md)",
    },
]


def build_agent_detail(doc, agent):
    add_heading(doc, agent["title"], level=2, color=COLOR_HEADING)

    for label, val in [("역할", agent["role"]), ("사용 시점", agent["when"]), ("사용 도구", agent["tools"])]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}  ")
        r.bold = True
        r.font.color.rgb = COLOR_TITLE
        p.add_run(val)

    p_area = doc.add_paragraph()
    r_area = p_area.add_run("담당 영역")
    r_area.bold = True
    r_area.font.color.rgb = COLOR_TITLE

    area_rows = [(name, desc) for name, desc in agent["areas"]]
    make_table(doc, ["영역", "설명"], area_rows, [Cm(4), Cm(10)], header_bg=COLOR_SUB_HDR)

    doc.add_paragraph()

    p_loc = doc.add_paragraph()
    r_loc = p_loc.add_run("자료 위치  ")
    r_loc.bold = True
    r_loc.font.color.rgb = COLOR_TITLE
    p_loc.add_run(agent["location"])

    add_note(doc, agent["note"])
    doc.add_paragraph()


def build_agents_section(doc):
    add_heading(doc, "2. 에이전트 상세 설명", level=1, color=COLOR_HEADING)
    doc.add_paragraph()
    for agent in AGENTS_DETAIL:
        build_agent_detail(doc, agent)
    doc.add_page_break()


# ── 협업 패턴 ─────────────────────────────────────────────

def build_collab_section(doc):
    add_heading(doc, "3. 에이전트 협업 패턴", level=1, color=COLOR_HEADING)

    p = doc.add_paragraph(
        "리더(오케스트레이터)가 요청을 도메인별 작업으로 분해하고, "
        "수집 팀원을 병렬 스폰한 뒤 결과를 writer 에게 전달합니다. "
        "팀원당 5~6개 작업, 전체 3~5명 규모를 권장합니다."
    )
    p.paragraph_format.space_after = Pt(8)

    headers = ["사용 케이스", "구성 팀원", "규모"]
    rows = [
        ("정형 보고서 (주간/월간/회의록)", "data-teammate + ops-teammate + document-writer",                                          "3명"),
        ("과제 진척률 심층 보고",          "data-teammate + ops-teammate + document-writer",                                          "3명"),
        ("신규 과제 브레인스토밍",          "tech-research-teammate + data-teammate + ops-teammate + document-writer",                 "4명"),
        ("경쟁사 분석 → KPI 재설정",       "tech-research-teammate + data-teammate (+ ops-teammate 선택) + document-writer", "3~4명"),
    ]
    col_widths = [Cm(5), Cm(10), Cm(2)]
    make_table(doc, headers, rows, col_widths)

    doc.add_paragraph()

    add_heading(doc, "팀 운영 SOP 요약", level=2, color=COLOR_HEADING)
    steps = [
        ("1. 분해",   "리더가 요청을 도메인별 작업으로 쪼갠다 (사내 수치 / 사내 운영 / 사외 리서치 / 작성)."),
        ("2. 병렬 수집", "수집 팀원(data · ops · tech-research)을 동시 스폰, 출처와 함께 데이터 수집."),
        ("3. 위임 작성", "결과를 모아 document-writer(.md) 또는 ppt-writer(.pptx)에 전달, placeholder 치환."),
        ("4. 품질 게이트", "TaskCompleted hook 이 outputs/*.md 의 출처 인용을 자동 검증."),
    ]
    for label, desc in steps:
        p = doc.add_paragraph()
        r_l = p.add_run(f"{label}  ")
        r_l.bold = True
        r_l.font.color.rgb = COLOR_ACCENT
        p.add_run(desc)

    doc.add_paragraph()
    add_note(doc, "(출처: CLAUDE.md 팀 운영 패턴(SOP) 섹션)")


# ── 메인 ──────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()
    set_font(doc)

    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    build_cover(doc)
    build_overview(doc)
    build_agents_section(doc)
    build_collab_section(doc)

    doc.save(OUTPUT_FILE)
    print(f"저장 완료: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
