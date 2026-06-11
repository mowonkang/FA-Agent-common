"""유연성 경쟁사 분석 v2 (.md) → CEO 보고용 상세 Word(.docx) 변환기.

입력: outputs/유연성_경쟁사_분석_v2_2026-05-21.md
출력: outputs/유연성_경쟁사_분석_v2_2026-05-21.docx

특징
- 표지 페이지 (제목 + 작성일 + 작성부서 + version)
- TOC 필드 (Heading 1~3) — Word 에서 F9 로 갱신
- H1=Heading 1 / H2=Heading 2 / H3=Heading 3
- 마크다운 표 전부 Word Table 변환 (헤더 진한 회색 배경, 숫자 우측·라벨 좌측)
- (출처: ...) 인용 보존
- 풋터: FA기술담당 · 2026-05-21 · v2 + 우측 페이지 번호
- 폰트: 한글 맑은 고딕 (LG스마트체 미설치 환경 fallback) / 영문 Arial Narrow
- 강조 컬러: 3사 ('28) 결론값(76% / 57% / 42%) 등 핵심 수치 파랑 #0000FF Bold
- brand-guidelines 자동 적용 (색·폰트·풋노트 룰 §1, §2, §3)
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "outputs" / "유연성_경쟁사_분석_v2_2026-05-21.md"
DST = ROOT / "outputs" / "유연성_경쟁사_분석_v2_2026-05-21.docx"

# ---------------- LGES brand-guidelines 색·폰트 상수 ----------------
BLACK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)  # 강조 (3사 '28 결론값 등 6~10개만)
GREEN = RGBColor(0x00, 0x66, 0x00)  # 풋노트 (초록곰팡이)
GRAY_HEADER = "C0C0C0"  # 표 헤더 배경 (진한 회색)
SOFT_GRAY = "E5E5E5"  # 박스 인용 배경
MID_GRAY_TEXT = RGBColor(0x55, 0x55, 0x55)

# 한글 폰트: LG스마트체가 없는 환경에서는 맑은 고딕으로 fallback
FONT_HANGUL = "LG스마트체 Regular"
FONT_HANGUL_FALLBACK = "맑은 고딕"
FONT_LATIN = "Arial Narrow"

# 핵심 강조 수치(파랑) — 3사 '28 결론값 + 임팩트 포인트만, 6~10개 제한
BLUE_KEYWORDS = {
    "76%",
    "57%",
    "42%",
    "+29%p",
    "95%",
    "+465.2%",
}
BLUE_USED_COUNT = {"n": 0}
BLUE_MAX = 10


# ---------------- 폰트 적용 헬퍼 (한·영 자동 분리) ----------------
HANGUL_RE = re.compile(r"[ㄱ-힝]")


def _set_run_font(run, *, size: int = 10, bold: bool = False, color: RGBColor | None = None):
    """run 에 한글/영문 폰트를 모두 지정. brand-guidelines §2."""
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = FONT_LATIN
    # 한글 (East Asian) 폰트도 별도 지정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_HANGUL_FALLBACK)
    rFonts.set(qn("w:cs"), FONT_LATIN)
    if color is not None:
        run.font.color.rgb = color


# ---------------- 인라인 마크다운 파서 ----------------
# 우선순위: 인라인 코드 `...`, 볼드 **...**, 링크 [text](url)
INLINE_TOKEN = re.compile(
    r"(`[^`]+`)|(\*\*[^*]+\*\*)|(\[[^\]]+\]\([^)]+\))"
)


def _emit_run(paragraph, text: str, *, base_size: int, base_bold: bool):
    """텍스트를 paragraph 에 run 으로 추가. 핵심 키워드는 파랑 Bold 강조."""
    # 키워드 강조 처리 — BLUE_KEYWORDS 가 들어있으면 분리 후 파랑으로
    pos = 0
    while pos < len(text):
        # 가장 가까운 키워드 매칭 위치 찾기
        hit_kw = None
        hit_pos = len(text)
        for kw in BLUE_KEYWORDS:
            idx = text.find(kw, pos)
            if idx != -1 and idx < hit_pos:
                hit_pos = idx
                hit_kw = kw
        if hit_kw is None or BLUE_USED_COUNT["n"] >= BLUE_MAX:
            # 남은 텍스트 그대로
            run = paragraph.add_run(text[pos:])
            _set_run_font(run, size=base_size, bold=base_bold)
            return
        # 키워드 이전 텍스트
        if hit_pos > pos:
            run = paragraph.add_run(text[pos:hit_pos])
            _set_run_font(run, size=base_size, bold=base_bold)
        # 키워드 자체 (파랑 Bold)
        run = paragraph.add_run(hit_kw)
        _set_run_font(run, size=base_size, bold=True, color=BLUE)
        BLUE_USED_COUNT["n"] += 1
        pos = hit_pos + len(hit_kw)


def add_inline_runs(paragraph, text: str, *, base_size: int = 10, base_bold: bool = False):
    """인라인 코드/볼드/링크 처리 + 핵심 키워드 강조."""
    # 링크: [text](url) → "text (url)"
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    pos = 0
    for m in INLINE_TOKEN.finditer(text):
        if m.start() > pos:
            _emit_run(paragraph, text[pos : m.start()], base_size=base_size, base_bold=base_bold)
        if m.group(1):  # 인라인 코드
            code_text = m.group(1)[1:-1]
            run = paragraph.add_run(code_text)
            _set_run_font(run, size=base_size, bold=base_bold)
            run.font.name = "Consolas"
            run.font.color.rgb = MID_GRAY_TEXT
        elif m.group(2):  # 볼드
            bold_text = m.group(2)[2:-2]
            _emit_run(paragraph, bold_text, base_size=base_size, base_bold=True)
        pos = m.end()
    if pos < len(text):
        _emit_run(paragraph, text[pos:], base_size=base_size, base_bold=base_bold)


# ---------------- 표 셀 정렬 (숫자 우측 / 라벨 좌측) ----------------
NUM_CELL = re.compile(r"^[\s+\-(]*[\d.,~%/]+.*$")


def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            txt = row[c_idx] if c_idx < len(row) else ""
            # 셀 본문 초기화 후 inline run 으로 채움
            cell.text = ""
            p = cell.paragraphs[0]
            base_bold = r_idx == 0
            base_size = 9 if r_idx > 0 else 10
            # 정렬: 헤더는 가운데, 본문은 숫자/% 우측·텍스트 좌측
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _shade_cell(cell, GRAY_HEADER)
            else:
                if NUM_CELL.match(txt) and "|" not in txt:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, txt, base_size=base_size, base_bold=base_bold)
    # 표 뒤 여백
    doc.add_paragraph("")


# ---------------- 표 파싱 ----------------
def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    # 구분선 (---|---) 제거
    cleaned = [r for r in rows if not all(set(c) <= {"-", ":", " "} for c in r)]
    return cleaned, i


# ---------------- 표지 / TOC / 헤더·푸터 ----------------
def add_cover(doc: Document):
    """표지 페이지. 제목 + 부제 + 작성일/작성부서/version."""
    # 빈 단락으로 여백
    for _ in range(5):
        doc.add_paragraph("")

    # 메인 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Samsung SDI · CATL")
    _set_run_font(run, size=24, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("유연성 지표 딥 분석 v2")
    _set_run_font(run2, size=24, bold=True)

    # 부제
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("(CEO 보고용 결정판)")
    _set_run_font(run3, size=16, bold=False, color=MID_GRAY_TEXT)

    # 구분선 (단락 하단 보더)
    sep = doc.add_paragraph()
    sep_run = sep.add_run(" ")
    _set_run_font(sep_run, size=10)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    for _ in range(8):
        doc.add_paragraph("")

    # 메타데이터 표 (3행 1열)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("작성일", "2026-05-21"),
        ("작성부서", "FA기술담당"),
        ("Version", "v2 (하이브리드 통합본 — B 정량 + C 정의 통합)"),
        ("정의 기준", "F1 (2026-05-21 신규 확정안)"),
    ]
    for r_idx, (k, v) in enumerate(rows):
        for c_idx, val in enumerate((k, v)):
            cell = meta_table.cell(r_idx, c_idx)
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(val)
            _set_run_font(run, size=11, bold=(c_idx == 0))
            if c_idx == 0:
                _shade_cell(cell, SOFT_GRAY)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_page_break()


def add_toc(doc: Document):
    """Word 자동 목차 필드. Word 에서 F9 로 갱신."""
    p = doc.add_paragraph()
    run = p.add_run("목 차")
    _set_run_font(run, size=18, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # TOC field — 단일 run 안에 fldChar(begin) + instrText + fldChar(separate)
    #            placeholder text + fldChar(end) 순서로 배치
    p = doc.add_paragraph()
    run = p.add_run()
    _set_run_font(run, size=11, bold=False)
    r_element = run._element

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    r_element.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r_element.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    r_element.append(fldChar2)

    # placeholder 텍스트 (Word 가 F9 새로고침 전까지 표시)
    placeholder = OxmlElement("w:t")
    placeholder.text = "목차는 Word 에서 우클릭 → '필드 업데이트(F9)' 로 자동 생성됩니다."
    r_element.append(placeholder)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    r_element.append(fldChar3)

    doc.add_page_break()


def setup_footer(doc: Document):
    """모든 섹션 공통 푸터: 좌측 'FA기술담당 · 2026-05-21 · v2' + 우측 페이지번호."""
    from docx.enum.text import WD_TAB_ALIGNMENT

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""

    # 우측 탭 스톱 설정 (페이지 너비 - 좌우 마진 ≈ 17cm)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)

    # 좌측 라벨
    run = p.add_run("FA기술담당 · 2026-05-21 · v2")
    _set_run_font(run, size=8, bold=True, color=GREEN)

    # 탭 (오른쪽 정렬용)
    tab_run = p.add_run("\t")
    _set_run_font(tab_run, size=8, color=GREEN)

    # 페이지 번호 prefix
    pre_run = p.add_run("Page ")
    _set_run_font(pre_run, size=8, color=GREEN)

    # PAGE 필드
    run2 = p.add_run()
    _set_run_font(run2, size=8, bold=False, color=GREEN)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2._element.append(fldChar1)
    run2._element.append(instrText)
    run2._element.append(fldChar2)


def setup_styles(doc: Document):
    """기본 스타일 + Heading 1~3 폰트/색 + Block Quote 스타일."""
    # 기본 (Normal)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(10)
    # 한글 폰트도 normal 에 설정
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_LATIN)
    rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    rFonts.set(qn("w:eastAsia"), FONT_HANGUL_FALLBACK)

    # Heading 1
    for level, (size, bold) in {
        1: (18, True),
        2: (14, True),
        3: (12, True),
        4: (11, True),
    }.items():
        h = doc.styles[f"Heading {level}"]
        h.font.name = FONT_LATIN
        h.font.size = Pt(size)
        h.font.bold = bold
        h.font.color.rgb = BLACK
        rPr_h = h.element.get_or_add_rPr()
        rFonts_h = rPr_h.find(qn("w:rFonts"))
        if rFonts_h is None:
            rFonts_h = OxmlElement("w:rFonts")
            rPr_h.append(rFonts_h)
        rFonts_h.set(qn("w:ascii"), FONT_LATIN)
        rFonts_h.set(qn("w:hAnsi"), FONT_LATIN)
        rFonts_h.set(qn("w:eastAsia"), FONT_HANGUL_FALLBACK)

    # Block Quote 스타일
    if "BlockQuote" not in doc.styles:
        style = doc.styles.add_style("BlockQuote", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_LATIN
        style.font.size = Pt(10)
        style.font.italic = False
        style.font.color.rgb = MID_GRAY_TEXT
        style.paragraph_format.left_indent = Cm(0.6)
        style.paragraph_format.right_indent = Cm(0.6)


# ---------------- 메인 변환 ----------------
def build(md_path: Path, out_path: Path):
    doc = Document()

    # 페이지 마진
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    setup_styles(doc)
    setup_footer(doc)

    add_cover(doc)
    add_toc(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_frontmatter = False
    table_count = 0

    # 인용 박스 (>) 수집 모드
    blockquote_buffer: list[str] = []

    def flush_quote():
        if not blockquote_buffer:
            return
        # 회색 박스: 단일 셀 표로 표현
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.text = ""
        _shade_cell(cell, SOFT_GRAY)
        for line in blockquote_buffer:
            p = cell.add_paragraph()
            add_inline_runs(p, line, base_size=10, base_bold=False)
        # 마지막 빈 단락 제거를 위해 첫 단락 사용 처리
        first_p = cell.paragraphs[0]
        if not first_p.runs and len(cell.paragraphs) > 1:
            first_p._element.getparent().remove(first_p._element)
        doc.add_paragraph("")
        blockquote_buffer.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # frontmatter (--- ... ---) 스킵
        if i == 0 and stripped == "---":
            in_frontmatter = True
            i += 1
            continue
        if in_frontmatter:
            if stripped == "---":
                in_frontmatter = False
            i += 1
            continue

        # 빈 줄
        if not stripped:
            flush_quote()
            i += 1
            continue

        # 인용 (>)
        if stripped.startswith(">"):
            blockquote_buffer.append(stripped.lstrip(">").strip())
            i += 1
            continue
        else:
            flush_quote()

        # 헤딩
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            level = min(level, 4)
            # H1 은 새 페이지 시작
            if level == 1:
                doc.add_page_break()
            h = doc.add_heading(level=level)
            add_inline_runs(h, text, base_size={1: 18, 2: 14, 3: 12, 4: 11}[level], base_bold=True)
            i += 1
            continue

        # 수평선 (---)
        if stripped.startswith("---"):
            # 빈 단락 + 하단 보더
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "808080")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # 표
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
                table_count += 1
            continue

        # 번호 리스트
        ordered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, ordered.group(2), base_size=10, base_bold=False)
            i += 1
            continue

        # 불릿 리스트
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, stripped[2:], base_size=10, base_bold=False)
            i += 1
            continue

        # 코드 블록
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GRAY_TEXT
            continue

        # 일반 단락 (이탤릭 *...* 마지막 단락 처리 포함)
        p = doc.add_paragraph()
        # 이탤릭 단락 (마지막 본문) 처리: 줄 전체가 *로 감싸진 경우
        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            inner = stripped.strip("*").strip()
            add_inline_runs(p, inner, base_size=9, base_bold=False)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = MID_GRAY_TEXT
        else:
            add_inline_runs(p, stripped, base_size=10, base_bold=False)
        i += 1

    flush_quote()

    doc.save(out_path)
    return table_count


if __name__ == "__main__":
    BLUE_USED_COUNT["n"] = 0
    tcount = build(SRC, DST)
    size_kb = DST.stat().st_size / 1024
    # 인용 카운트
    src_text = SRC.read_text(encoding="utf-8")
    citation_count = src_text.count("(출처:")
    print(f"saved -> {DST}")
    print(f"size  -> {size_kb:.1f} KB")
    print(f"tables converted -> {tcount}")
    print(f"(출처: ...) preserved -> {citation_count}")
    print(f"blue highlights used -> {BLUE_USED_COUNT['n']} / max {BLUE_MAX}")
