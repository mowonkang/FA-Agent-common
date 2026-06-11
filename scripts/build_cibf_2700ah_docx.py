"""CIBF 2026 2700Ah 종합보고서 md -> docx 변환기.

소스: references/기술자료/2026-05-26_CIBF2026_2700Ah셀_고중량핸들링자동화_종합보고서.md
출력: outputs/CIBF2026_2700Ah셀_고중량핸들링자동화_종합보고_2026-05-26.docx

YAML frontmatter 를 건너뛰고, 헤더/표/리스트/인용/구분선을 Word 로 변환한다.
LGES 사내 톤: 한글 본문 맑은 고딕, 표 헤더 볼드, 풋노트/주의 회색.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "references" / "기술자료" / "2026-05-26_CIBF2026_2700Ah셀_고중량핸들링자동화_종합보고서.md"
DST = ROOT / "outputs" / "CIBF2026_2700Ah셀_고중량핸들링자동화_종합보고_2026-05-26.docx"

KO_FONT = "맑은 고딕"
EN_FONT = "Arial"

BOLD = re.compile(r"\*\*([^*]+)\*\*")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_CODE = re.compile(r"`([^`]+)`")


def _set_ko_font(run):
    run.font.name = EN_FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    if rFonts is None:
        from docx.oxml.ns import qn

        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    from docx.oxml.ns import qn

    rFonts.set(qn("w:eastAsia"), KO_FONT)


def add_runs(paragraph, text: str) -> None:
    text = LINK.sub(r"\1 (\2)", text)
    pattern = re.compile(r"`([^`]+)`|\*\*([^*]+)\*\*")
    i = 0
    pieces = []
    for m in pattern.finditer(text):
        if m.start() > i:
            pieces.append(("plain", text[i:m.start()]))
        if m.group(1) is not None:
            pieces.append(("code", m.group(1)))
        else:
            pieces.append(("bold", m.group(2)))
        i = m.end()
    if i < len(text):
        pieces.append(("plain", text[i:]))
    if not pieces:
        pieces = [("plain", text)]
    for kind, value in pieces:
        run = paragraph.add_run(value)
        _set_ko_font(run)
        if kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB1, 0x29, 0x4A)
        elif kind == "bold":
            run.bold = True


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    cleaned = [r for r in rows if not all(set(c) <= {"-", ":", " "} for c in r)]
    return cleaned, i


def strip_frontmatter(lines):
    if lines and lines[0].strip() == "---":
        for j in range(1, len(lines)):
            if lines[j].strip() == "---":
                return lines[j + 1:]
    return lines


def build(md_path: Path, out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = Pt(10.5)

    lines = strip_frontmatter(md_path.read_text(encoding="utf-8").splitlines())
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, text)
            i += 1
            continue

        if set(stripped) <= {"-"} and len(stripped) >= 3:
            p = doc.add_paragraph()
            run = p.add_run("―" * 40)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            table.autofit = True
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs(p, row[c_idx] if c_idx < len(row) else "")
                    for run in p.runs:
                        run.font.size = Pt(8.5)
                        if r_idx == 0:
                            run.bold = True
            doc.add_paragraph("")
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, ordered.group(2))
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            indent = len(lines[i]) - len(lines[i].lstrip())
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 2:
                p.paragraph_format.left_indent = Pt(18 + indent * 4)
            add_runs(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    # 표지 성격의 부제/날짜 한 줄을 맨 앞에 추가
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"saved -> {out_path}")


if __name__ == "__main__":
    build(SRC, DST)
