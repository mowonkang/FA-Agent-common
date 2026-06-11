"""md -> docx 변환 (Agent Teams 가이드 전용 간이 변환기).

GitHub Flavored Markdown 의 모든 요소를 다 지원하지는 않지만,
이 문서에서 사용된 헤더/단락/리스트/표/코드블록/인라인 강조/링크/인용을 지원한다.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "outputs" / "Claude_Agent_Teams_가이드.md"
DST = ROOT / "outputs" / "Claude_Agent_Teams_가이드.docx"


def add_code_style(doc: Document) -> None:
    styles = doc.styles
    if "CodeBlock" in styles:
        return
    style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text: str) -> None:
    """인라인 코드/볼드/링크 처리."""
    text = LINK.sub(r"\1 (\2)", text)
    pieces = []
    i = 0
    pattern = re.compile(r"`([^`]+)`|\*\*([^*]+)\*\*")
    for m in pattern.finditer(text):
        if m.start() > i:
            pieces.append(("plain", text[i : m.start()]))
        if m.group(1) is not None:
            pieces.append(("code", m.group(1)))
        else:
            pieces.append(("bold", m.group(2)))
        i = m.end()
    if i < len(text):
        pieces.append(("plain", text[i:]))

    for kind, value in pieces:
        run = paragraph.add_run(value)
        if kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xB1, 0x29, 0x4A)
        elif kind == "bold":
            run.bold = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    cleaned = [r for r in rows if not all(set(c) <= {"-", ":", " "} for c in r)]
    return cleaned, i


def build(md_path: Path, out_path: Path) -> None:
    doc = Document()
    add_code_style(doc)

    body_style = doc.styles["Normal"]
    body_style.font.name = "맑은 고딕"
    body_style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph(style="CodeBlock")
            run = p.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        if stripped.startswith("---"):
            doc.add_paragraph("―" * 30)
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_runs(p, row[c_idx] if c_idx < len(row) else "")
                    if r_idx == 0:
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph("")
            continue

        ordered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if ordered:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, ordered.group(2))
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        # 일반 단락
        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"saved -> {out_path}")


if __name__ == "__main__":
    build(SRC, DST)
